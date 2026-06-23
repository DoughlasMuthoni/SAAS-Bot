import re

from app.ingestion.processors.base import BaseProcessor, ProcessedDocument


class UploadProcessor(BaseProcessor):
    async def process(self, metadata_json: dict, file_path: str | None = None) -> list[ProcessedDocument]:
        if not file_path:
            return []
        mime_type = (metadata_json or {}).get("mime_type", "")
        name = (metadata_json or {}).get("name", "Uploaded File")

        if "pdf" in mime_type:
            return await self._process_pdf(file_path, name)
        elif mime_type in ("text/plain", "text/csv"):
            return await self._process_text(file_path, name, mime_type)
        elif "wordprocessingml" in mime_type or "docx" in mime_type:
            return await self._process_docx(file_path, name)
        return []

    async def _process_pdf(self, path: str, name: str) -> list[ProcessedDocument]:
        import fitz  # PyMuPDF
        docs = []
        try:
            pdf = fitz.open(path)
            for page_num, page in enumerate(pdf, start=1):
                text = page.get_text()
                if text.strip():
                    normalized = re.sub(r"\s+", " ", text).strip()
                    docs.append(ProcessedDocument(
                        title=f"{name} — Page {page_num}",
                        raw_content=text,
                        normalized_content=normalized,
                        doc_type="pdf",
                        page_number=page_num,
                    ))
            pdf.close()
        except Exception as e:
            raise RuntimeError(f"PDF extraction failed: {e}") from e
        return docs

    async def _process_text(self, path: str, name: str, mime_type: str) -> list[ProcessedDocument]:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        normalized = re.sub(r"\s+", " ", content).strip()
        return [ProcessedDocument(
            title=name,
            raw_content=content,
            normalized_content=normalized,
            doc_type="csv" if "csv" in mime_type else "text",
        )]

    async def _process_docx(self, path: str, name: str) -> list[ProcessedDocument]:
        from docx import Document
        doc = Document(path)
        blocks: list[str] = []

        # Paragraphs
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                blocks.append(text)

        # Tables — render each row as "Col1: val | Col2: val"
        for table in doc.tables:
            headers: list[str] = []
            for i, row in enumerate(table.rows):
                cells = [c.text.strip() for c in row.cells]
                if i == 0:
                    headers = cells
                    blocks.append(" | ".join(cells))
                else:
                    if headers:
                        row_text = " | ".join(
                            f"{headers[j]}: {v}" if j < len(headers) else v
                            for j, v in enumerate(cells)
                        )
                    else:
                        row_text = " | ".join(cells)
                    if any(c for c in cells):
                        blocks.append(row_text)

        content = "\n\n".join(blocks)
        normalized = re.sub(r"\n{3,}", "\n\n", re.sub(r"[ \t]+", " ", content)).strip()
        return [ProcessedDocument(
            title=name,
            raw_content=content,
            normalized_content=normalized,
            doc_type="docx",
        )]
